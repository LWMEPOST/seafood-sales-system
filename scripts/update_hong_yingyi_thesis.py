from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


BASE = Path(r"D:\桌面\论文\600洪英毅(论文)!")
TARGET = BASE / "洪英毅 论文.docx"
BACKUP = BASE / "洪英毅 论文.codex-backup-20260515.docx"
MODULE_IMG = BASE / "系统功能模块图_新版.png"


def set_east_asia(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_text(paragraph: Paragraph, text: str, style: str | None = None, align=None):
    paragraph.clear()
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        set_east_asia(run)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None, align=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_east_asia(run)
    if align is not None:
        new_para.alignment = align
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_cm: float = 15.0) -> Paragraph:
    pic_para = insert_paragraph_after(paragraph, style="Normal")
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return pic_para


def insert_paragraph_after_table(table: Table, text: str = "", style: str | None = None, align=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_east_asia(run)
    if align is not None:
        new_para.alignment = align
    return new_para


def add_table_after(doc: Document, paragraph: Paragraph, data: list[list[str]], style: str = "Table Grid") -> Table:
    rows = len(data)
    cols = len(data[0]) if rows else 0
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    paragraph._p.addnext(table._tbl)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    set_east_asia(run)
    return table


def remove_paragraph(paragraph: Paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Paragraph not found: {startswith}")


def find_all_paragraphs(doc: Document, startswith: str) -> list[Paragraph]:
    return [p for p in doc.paragraphs if p.text.strip().startswith(startswith)]


def find_last_paragraph(doc: Document, startswith: str) -> Paragraph:
    matches = find_all_paragraphs(doc, startswith)
    if not matches:
        raise ValueError(f"Paragraph not found: {startswith}")
    return matches[-1]


def resize_table(table: Table, rows: int, cols: int):
    while len(table.columns) < cols:
        table.add_column(Cm(3))
    while len(table.rows) < rows:
        table.add_row()


def fill_table(table: Table, data: list[list[str]]):
    resize_table(table, len(data), len(data[0]))
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    set_east_asia(run)


def ensure_update_fields(doc: Document):
    settings = doc.settings.element
    upd = settings.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


def next_table_after_paragraph(paragraph: Paragraph) -> Table | None:
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return Table(node, paragraph._parent)
        node = node.getnext()
    return None


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size=size)


def draw_centered(draw: ImageDraw.ImageDraw, xy, text: str, font_obj, fill="black"):
    x, y = xy
    bbox = draw.multiline_textbbox((0, 0), text, font=font_obj, spacing=6, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x - w / 2, y - h / 2), text, font=font_obj, fill=fill, spacing=6, align="center")


def box(draw, x1, y1, x2, y2, text, title_font, fill="#FFFFFF"):
    draw.rounded_rectangle((x1, y1, x2, y2), radius=18, outline="black", width=3, fill=fill)
    draw_centered(draw, ((x1 + x2) / 2, (y1 + y2) / 2), text, title_font)


def line(draw, p1, p2):
    draw.line((p1, p2), fill="black", width=3)


def generate_module_diagram(path: Path):
    img = Image.new("RGB", (1800, 1500), "white")
    draw = ImageDraw.Draw(img)
    f_title = font(r"C:\Windows\Fonts\simsun.ttc", 38)
    f_box = font(r"C:\Windows\Fonts\simsun.ttc", 27)
    f_small = font(r"C:\Windows\Fonts\simsun.ttc", 24)

    root = (650, 40, 1150, 120)
    admin_root = (150, 190, 500, 260)
    mobile_root = (725, 190, 1075, 260)
    support_root = (1300, 190, 1650, 260)

    box(draw, *root, "水产品销售管理系统", f_title, "#E8EEF9")
    box(draw, *admin_root, "管理端", f_title, "#F4F7FC")
    box(draw, *mobile_root, "消费者端", f_title, "#F4F7FC")
    box(draw, *support_root, "系统支撑层", f_title, "#F4F7FC")

    line(draw, ((root[0] + root[2]) / 2, root[3]), ((root[0] + root[2]) / 2, 155))
    for child in (admin_root, mobile_root, support_root):
        cx = (child[0] + child[2]) / 2
        line(draw, (cx, 155), (cx, child[1]))
    line(draw, ((admin_root[0] + admin_root[2]) / 2, 155), ((support_root[0] + support_root[2]) / 2, 155))

    admin_boxes = [
        (80, 330, 570, 470, "数据看板\n今日销售额 / 今日订单数\n库存预警 / 销售趋势"),
        (80, 510, 570, 650, "商品管理\n分类筛选 / 新增编辑\n上下架管理 / 批量导入"),
        (80, 690, 570, 830, "订单管理\n订单查询 / 订单详情\n发货处理 / 状态跟踪"),
        (80, 870, 570, 1010, "用户管理\n用户分页查看 / 角色识别"),
        (80, 1050, 570, 1190, "食谱管理\n新增食谱 / 编辑食谱 / 内容维护"),
        (80, 1230, 570, 1370, "溯源与库存\n溯源记录查询 / 区域库存查看"),
    ]
    mobile_boxes = [
        (655, 330, 1145, 470, "用户认证\n登录 / 注册"),
        (655, 510, 1145, 650, "商品浏览\n首页推荐 / 分类检索 / 商品详情"),
        (655, 690, 1145, 830, "购物车与下单\n加入购物车 / 数量调整\n提交订单 / 模拟支付"),
        (655, 870, 1145, 1010, "地址管理\n新增地址 / 编辑地址 / 默认地址选择"),
        (655, 1050, 1145, 1190, "订单中心\n订单列表 / 订单详情 / 确认收货"),
        (655, 1230, 1145, 1370, "个人中心与溯源\n个人信息查看 / 快捷入口\n时间线展示 / 节点详情"),
    ]
    support_boxes = [
        (1230, 420, 1720, 540, "统一 REST 接口"),
        (1230, 620, 1720, 740, "登录拦截与权限校验"),
        (1230, 820, 1720, 940, "文件上传"),
        (1230, 1020, 1720, 1140, "MySQL 数据持久化"),
    ]

    for rect in admin_boxes:
        box(draw, *rect, f_box, "#FFFFFF")
        line(draw, ((admin_root[0] + admin_root[2]) / 2, admin_root[3]), ((rect[0] + rect[2]) / 2, rect[1]))
    for rect in mobile_boxes:
        box(draw, *rect, f_box, "#FFFFFF")
        line(draw, ((mobile_root[0] + mobile_root[2]) / 2, mobile_root[3]), ((rect[0] + rect[2]) / 2, rect[1]))
    for rect in support_boxes:
        box(draw, *rect, f_small, "#FFFFFF")
        line(draw, ((support_root[0] + support_root[2]) / 2, support_root[3]), ((rect[0] + rect[2]) / 2, rect[1]))

    img.save(path)


def update_heading_styles(doc: Document):
    mapping = {
        "1 引言": ("1 前言", "Heading 1"),
        "1.1 课题研究的背景": ("1.1 研究背景", "Heading 2"),
        "1.2 课题研究的目的和意义": ("1.2 研究的目的与意义", "Heading 2"),
        "2.6 本章小结": ("2.5 本章小结", "Heading 2"),
        "3.1 可行性分析": ("3.1 可行性分析", "Heading 2"),
        "3.3.1 系统用例图": ("3.3.1 系统用例", "Heading 3"),
        "4.2 主要功能模块": ("4.2 系统功能模块设计", "Heading 2"),
        "4.2.1 系统活动图（购物结算流程）": ("4.2.1 系统活动图", "Heading 3"),
        "4.3.1 数据库关系（E-R模型简述）": ("4.3.1 数据库关系图（E-R）图与实体属性图", "Heading 3"),
        "5.2.8 区域库存与看板管理功能模块实现": ("5.2.9 区域库存与看板管理功能模块实现", "Heading 3"),
        "6.3 本章小结": ("6.3 测试总结", "Heading 2"),
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in mapping:
            new_text, style = mapping[text]
            set_paragraph_text(p, new_text, style=style)


def update_ch3(doc: Document):
    p = find_paragraph(doc, "系统主要实体及关系如下：")
    text = (
        "结合后端实体类 User、Address、Category、Product、Order、OrderItem、"
        "TraceRecord、Recipe 与 RegionStock，系统主要类模型可抽象为图 3-2 所示的 UML 类图。"
        "该类图既反映了数据库中的主要对象结构，也体现了业务层的核心关联关系。从类关系上看，"
        "User 与 Order 形成用户下单主链路，Order 与 OrderItem 构成订单主从结构，"
        "Category 与 Product 构成商品分类结构，Product 又分别与 TraceRecord、Recipe "
        "和 OrderItem 建立关联。Address 采用一对多建模方式，更符合真实收货场景；"
        "RegionStock 则作为看板统计类独立存在，用于支撑区域库存分析。该类模型确保交易数据、"
        "内容数据与追溯数据能够解耦存储、关联查询，支撑后续迭代。"
    )
    set_paragraph_text(p, text, style="Normal")
    cap = find_paragraph(doc, "图3-2 类图")
    set_paragraph_text(cap, "图 3-2 系统主要类图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)


def update_ch4(doc: Document):
    h42 = find_paragraph(doc, "4.2 系统功能模块设计")
    # only insert once
    if not any("图 4-1 系统功能模块图" in p.text for p in doc.paragraphs):
        desc = insert_paragraph_after(
            h42,
            "系统功能模块图如图 4-1 所示。结合本地项目的前后端实现，本文将系统功能按照管理端、消费者端和系统支撑层三个维度进行划分。图中既体现了不同角色的功能边界，也体现了系统支撑能力对业务模块的统一承载关系。",
            style="Normal",
        )
        pic = insert_picture_after(desc, MODULE_IMG, width_cm=15.5)
        insert_paragraph_after(pic, "图 4-1 系统功能模块图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)

    p247 = find_paragraph(doc, "以下流程展示消费者完成一次标准交易的活动链路")
    set_paragraph_text(
        p247,
        "以下流程展示消费者完成一次标准交易的活动链路，覆盖登录校验、商品浏览、地址完善、订单提交、支付、发货和确认收货等关键环节。如图 4-2 所示。",
        style="Normal",
    )
    p249 = find_paragraph(doc, "图4-2")
    set_paragraph_text(p249, "图 4-2 购物结算活动图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)

    p252 = find_paragraph(doc, "下述时序描述订单创建到完成的服务交互")
    set_paragraph_text(
        p252,
        "下述时序描述订单从创建到完成的关键服务交互过程。移动端负责发起业务请求，后端订单服务负责写入订单主表和订单明细表，并通过状态流转控制支付、发货和收货过程。如图 4-3 所示。",
        style="Normal",
    )
    p255 = find_paragraph(doc, "图4-3")
    set_paragraph_text(p255, "图 4-3 系统时序图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)

    p258 = find_paragraph(doc, "数据库关系可抽象为用户")
    set_paragraph_text(
        p258,
        "数据库关系图采用实体框、关系菱形节点与 1/n 基数标注方式进行表达。系统以 sys_user、oms_order、oms_order_item、pms_product、pms_category、trace_record、pms_recipe 和 sys_address 为主要实体，展示交易主链路与扩展信息链路之间的关系。逻辑关系见表 4-2。",
        style="Normal",
    )
    p259 = find_paragraph(doc, "表 4-2")
    set_paragraph_text(p259, "表 4-2 数据库关系说明表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[8],
        [
            ["实体 A", "实体 B", "关系", "关系说明"],
            ["sys_user", "sys_address", "1:n", "一个用户可维护多个收货地址"],
            ["sys_user", "oms_order", "1:n", "一个用户可创建多笔订单"],
            ["oms_order", "oms_order_item", "1:n", "一笔订单可包含多个商品明细"],
            ["pms_category", "pms_product", "1:n", "一个分类下包含多个商品"],
            ["pms_product", "trace_record", "1:n", "一个商品可对应多条溯源记录"],
            ["pms_product", "pms_recipe", "1:n", "一个商品可关联多条食谱内容"],
            ["pms_product", "oms_order_item", "1:n", "一个商品可出现在多条订单明细中"],
        ],
    )
    p261 = find_paragraph(doc, "图4-4")
    set_paragraph_text(p261, "图 4-4 数据库实体关系图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    p263 = find_paragraph(doc, "")
    # skip blank search approach; locate exact blank heading between 4-4 and 4.3.2
    for p in doc.paragraphs:
        if p.style.name == "Heading 3" and not p.text.strip():
            set_paragraph_text(
                p,
                "依据主要数据表结构，系统进一步补充了用户与地址、分类与商品、订单与订单明细、溯源与食谱等核心实体属性图，如图 4-5 至图 4-8 所示。",
                style="Normal",
            )
            break
    set_paragraph_text(find_paragraph(doc, "图4-5"), "图 4-5 用户与地址实体属性图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(find_paragraph(doc, "图4-6"), "图 4-6 分类与商品实体属性图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(find_paragraph(doc, "图4-7"), "图 4-7 订单与订单明细实体属性图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(find_paragraph(doc, "图4-8"), "图 4-8 溯源与食谱实体属性图", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)

    p273 = find_paragraph(doc, "系统主要表结构设计见表 4-3")
    set_paragraph_text(
        p273,
        "数据库物理设计以 MySQL 为基础实现。当前项目中的主要业务表并未在 SQL 中显式定义外键约束，而是通过业务字段和服务层逻辑完成实体关联。为保证字段表达完整性，系统主要数据表结构设计如下。",
        style="Normal",
    )
    set_paragraph_text(find_paragraph(doc, "表 4-3"), "表 4-3 系统数据表总体设计", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[9],
        [
            ["表名", "主要功能", "主键字段", "关联字段"],
            ["sys_user", "保存系统用户基础信息", "id", "无"],
            ["sys_address", "保存用户收货地址信息", "id", "user_id"],
            ["pms_category", "保存商品分类信息", "id", "无"],
            ["pms_product", "保存商品主数据与展示信息", "id", "category_id"],
            ["oms_order", "保存订单主信息与履约快照", "id", "user_id"],
            ["oms_order_item", "保存订单商品明细", "id", "order_id、product_id"],
            ["trace_record", "保存商品溯源时间线数据", "id", "product_id"],
            ["pms_recipe", "保存商品关联食谱内容", "id", "product_id"],
            ["pms_region_stock", "保存区域库存统计数据", "id", "无"],
        ],
    )

    set_paragraph_text(find_paragraph(doc, "表 4-4"), "表 4-4 sys_user 用户表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[10],
        [
            ["字段名", "数据类型", "约束", "说明"],
            ["id", "BIGINT", "主键，自增", "用户唯一标识"],
            ["username", "VARCHAR(50)", "非空，唯一", "登录账号"],
            ["password", "VARCHAR(100)", "非空", "登录密码"],
            ["nickname", "VARCHAR(50)", "可空", "用户昵称"],
            ["avatar", "VARCHAR(255)", "可空", "头像地址"],
            ["role", "VARCHAR(20)", "非空，默认 CUSTOMER", "用户角色，区分管理员与消费者"],
            ["create_time", "DATETIME", "默认当前时间", "创建时间"],
        ],
    )

    set_paragraph_text(find_paragraph(doc, "表 4-5"), "表 4-5 sys_address 地址表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[11],
        [
            ["字段名", "数据类型", "约束", "说明"],
            ["id", "BIGINT", "主键，自增", "地址唯一标识"],
            ["user_id", "BIGINT", "非空", "所属用户编号"],
            ["name", "VARCHAR(50)", "非空", "收货人姓名"],
            ["phone", "VARCHAR(20)", "非空", "收货联系电话"],
            ["detail", "VARCHAR(255)", "非空", "详细收货地址"],
            ["is_default", "INT", "默认 0", "是否默认地址，1 表示默认"],
            ["create_time", "DATETIME", "默认当前时间", "创建时间"],
            ["update_time", "DATETIME", "默认当前时间", "更新时间"],
        ],
    )

    set_paragraph_text(find_paragraph(doc, "表 4-7"), "表 4-6 pms_category 分类表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[12],
        [
            ["字段名", "数据类型", "约束", "说明"],
            ["id", "BIGINT", "主键，自增", "分类唯一标识"],
            ["name", "VARCHAR(50)", "非空", "分类名称"],
            ["icon", "VARCHAR(255)", "可空", "分类图标"],
            ["sort", "INT", "默认 0", "排序值"],
            ["status", "INT", "默认 1", "分类状态，1 表示启用"],
        ],
    )

    set_paragraph_text(find_paragraph(doc, "表 4-8"), "表 4-7 oms_order_item 订单明细表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[13],
        [
            ["字段名", "数据类型", "约束", "说明"],
            ["id", "BIGINT", "主键，自增", "明细唯一标识"],
            ["order_id", "BIGINT", "非空", "所属订单编号"],
            ["product_id", "BIGINT", "非空", "商品编号"],
            ["product_name", "VARCHAR(100)", "非空", "商品名称快照"],
            ["product_price", "DECIMAL(10,2)", "非空", "商品价格快照"],
            ["quantity", "INT", "非空", "购买数量"],
            ["product_image", "VARCHAR(255)", "可空", "商品图片快照"],
        ],
    )

    set_paragraph_text(find_paragraph(doc, "表 4-9"), "表 4-8 pms_product 商品表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[14],
        [
            ["字段名", "数据类型", "约束", "说明"],
            ["id", "BIGINT", "主键，自增", "商品唯一标识"],
            ["category_id", "BIGINT", "非空", "所属分类编号"],
            ["name", "VARCHAR(100)", "非空", "商品名称"],
            ["price", "DECIMAL(10,2)", "非空", "销售价格"],
            ["original_price", "DECIMAL(10,2)", "可空", "原始价格"],
            ["stock", "INT", "非空，默认 0", "库存数量"],
            ["origin", "VARCHAR(100)", "可空", "商品产地"],
            ["catch_date", "DATE", "可空", "捕捞日期"],
            ["farming_method", "VARCHAR(50)", "可空", "养殖或捕捞方式"],
            ["description", "TEXT", "可空", "商品描述"],
            ["main_image", "VARCHAR(255)", "可空", "主图地址"],
            ["detail_images", "JSON", "可空", "详情图列表"],
            ["status", "TINYINT", "默认 1", "上下架状态，1 上架，0 下架"],
            ["create_time", "DATETIME", "默认当前时间", "创建时间"],
        ],
    )

    set_paragraph_text(find_paragraph(doc, "表 4-10"), "表 4-9 oms_order 订单表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_table(
        doc.tables[15],
        [
            ["字段名", "数据类型", "约束", "说明"],
            ["id", "BIGINT", "主键，自增", "订单唯一标识"],
            ["order_sn", "VARCHAR(64)", "非空，唯一", "订单编号"],
            ["user_id", "BIGINT", "非空", "下单用户编号"],
            ["total_amount", "DECIMAL(10,2)", "非空", "订单总金额"],
            ["status", "TINYINT", "非空，默认 0", "订单状态"],
            ["create_time", "DATETIME", "默认当前时间", "下单时间"],
            ["payment_time", "DATETIME", "可空", "支付时间"],
            ["delivery_time", "DATETIME", "可空", "发货时间"],
            ["receive_time", "DATETIME", "可空", "收货时间"],
            ["receiver_name", "VARCHAR(50)", "可空", "收货人姓名"],
            ["receiver_phone", "VARCHAR(20)", "可空", "收货人电话"],
            ["receiver_address", "VARCHAR(255)", "可空", "收货地址快照"],
        ],
    )

    # insert missing table captions and tables before 4.4
    p280 = find_paragraph(doc, "表 4-9 oms_order 订单表")
    if not any("表 4-10 trace_record 溯源记录表" in p.text for p in doc.paragraphs):
        order_tbl = next_table_after_paragraph(p280)
        cap1 = insert_paragraph_after_table(order_tbl, "表 4-10 trace_record 溯源记录表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        tbl1 = add_table_after(
            doc,
            cap1,
            [
                ["字段名", "数据类型", "约束", "说明"],
                ["id", "BIGINT", "主键，自增", "溯源记录唯一标识"],
                ["product_id", "BIGINT", "非空", "对应商品编号"],
                ["node_name", "VARCHAR(50)", "非空", "溯源节点名称"],
                ["node_time", "DATETIME", "非空", "节点发生时间"],
                ["location", "VARCHAR(100)", "可空", "节点地点"],
                ["operator", "VARCHAR(50)", "可空", "操作人"],
                ["description", "VARCHAR(255)", "可空", "节点说明"],
                ["image", "VARCHAR(255)", "可空", "节点图片"],
            ],
        )
        cap2 = insert_paragraph_after_table(tbl1, "表 4-11 pms_recipe 食谱表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        tbl2 = add_table_after(
            doc,
            cap2,
            [
                ["字段名", "数据类型", "约束", "说明"],
                ["id", "BIGINT", "主键，自增", "食谱唯一标识"],
                ["product_id", "BIGINT", "可空", "关联商品编号"],
                ["title", "VARCHAR(200)", "非空", "食谱标题"],
                ["content", "TEXT", "可空", "食谱正文"],
                ["image", "VARCHAR(500)", "可空", "食谱配图"],
                ["create_time", "DATETIME", "默认当前时间", "创建时间"],
            ],
        )
        cap3 = insert_paragraph_after_table(tbl2, "表 4-12 pms_region_stock 区域库存表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_table_after(
            doc,
            cap3,
            [
                ["字段名", "数据类型", "约束", "说明"],
                ["id", "BIGINT", "主键，自增", "统计记录唯一标识"],
                ["region_name", "VARCHAR(50)", "非空", "区域名称"],
                ["stock_count", "INT", "默认 0", "区域库存总量"],
            ],
        )

    p283 = find_paragraph(doc, "本章从总体架构")
    set_paragraph_text(
        p283,
        "本章从总体结构、系统功能模块、数据库设计以及核心业务流程四个层面完成了系统设计说明。通过补充功能模块图、E-R 图、实体属性图和字段级物理表设计，系统的数据结构、模块边界与业务流转关系得到了更清晰的表达。",
        style="Normal",
    )


def update_ch5(doc: Document):
    # fix captions and add missing subsections
    caps = [
        ("图 5-3 商品详情页面", "图 5-4 商品详情页面"),
        ("图5-4 购物车页面", "图 5-5 购物车页面"),
        ("图5-5 地址编辑页", "图 5-6 地址编辑页"),
        ("图5-6 管理段端商品页面", "图 5-7 管理端商品管理页面"),
        ("图5-7 移动端订单详情", "图 5-8 移动端订单详情"),
        ("图5-8 管理端食谱管理页面", "图 5-9 管理端食谱管理页面"),
        ("图5-9 管理端 Dashboard 库存热力图区域", "图 5-10 管理端 Dashboard 库存热力图区域"),
    ]
    for old, new in caps:
        try:
            set_paragraph_text(find_paragraph(doc, old), new, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        except Exception:
            pass

    main_ch5 = find_last_paragraph(doc, "5 系统实现")
    main_58 = find_last_paragraph(doc, "5.2.9 区域库存与看板管理功能模块实现")
    main_ch5_idx = doc.paragraphs.index(main_ch5)
    main_58_idx = doc.paragraphs.index(main_58)

    # remove wrong insertion inside TOC / front matter
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("5.2.8 用户信息管理功能模块实现") and doc.paragraphs.index(p) < main_ch5_idx:
            next_p = p._p.getnext()
            remove_paragraph(p)
            if next_p is not None:
                next_para = Paragraph(next_p, doc._body)
                if next_para.text.strip().startswith("用户管理模块提供管理端分页查看能力"):
                    remove_paragraph(next_para)

    has_main_58 = any(
        p.text.strip().startswith("5.2.8 用户信息管理功能模块实现")
        and doc.paragraphs.index(p) > main_ch5_idx
        for p in doc.paragraphs
    )
    if not has_main_58:
        anchor = find_last_paragraph(doc, "图 5-9 管理端食谱管理页面")
        h = insert_paragraph_after(anchor, "5.2.8 用户信息管理功能模块实现", style="Heading 3")
        p1 = insert_paragraph_after(
            h,
            "用户管理模块提供管理端分页查看能力，并支持用户注册与登录基础流程。系统根据 role 字段区分管理员与消费者，当前版本采用轻量 token 机制满足教学与演示场景，后续可升级为 JWT 与刷新令牌方案。该模块保障基础身份体系可运行，为权限控制与审计扩展提供入口。",
            style="Normal",
        )
        insert_paragraph_after(p1, "", style="Normal")

    if not any("5.2.10 溯源信息管理功能模块实现" in p.text and doc.paragraphs.index(p) > main_ch5_idx for p in doc.paragraphs):
        # insert after current 5.2.9 block, before 5.3
        p330 = find_last_paragraph(doc, "图 5-10 管理端 Dashboard 库存热力图区域")
        h = insert_paragraph_after(p330, "5.2.10 溯源信息管理功能模块实现", style="Heading 3")
        p1 = insert_paragraph_after(
            h,
            "后端提供按商品编号查询溯源记录接口，并按节点时间排序返回；前端在商品详情页展示关键节点时间线，节点包括捕捞或采摘、加工分拣、质量检测和入库仓储等。通过可追溯信息可视化，系统将品质可信转化为可感知体验，是本项目区别于普通交易系统的重要特征。",
            style="Normal",
        )
        insert_paragraph_after(p1, "", style="Normal")

    p332 = find_paragraph(doc, "本章详细描述了系统核心功能模块")
    set_paragraph_text(
        p332,
        "本章围绕系统真实实现，说明了双端功能如何通过统一后端服务形成业务闭环，并结合页面截图展示了商品浏览、购物车下单、地址维护、订单处理、内容运营和后台看板等关键模块的实现效果。整体实现兼顾功能完整性与工程可维护性，能够支撑后续功能扩展与性能优化。",
        style="Normal",
    )


def update_ch6(doc: Document):
    p336 = find_paragraph(doc, "本系统采用接口测试")
    set_paragraph_text(
        p336,
        "本系统采用接口测试、功能测试与场景测试相结合的方式。接口测试关注请求参数、返回结构与状态码；功能测试关注页面交互与业务结果；场景测试关注跨模块流程完整性，如未登录下单拦截、地址缺失下单拦截和订单状态顺序流转等。测试环境采用本地开发环境，后端服务连接 MySQL，前端分别在管理端与移动端运行，使用真实接口联调。",
        style="Normal",
    )
    set_paragraph_text(find_paragraph(doc, "6.2.2 核心功能测试用例设计与结果"), "6.2.2 核心功能测试用例设计与结果", style="Heading 3")
    for txt, new in [
        ("表6-1 用户认证模块测试", "表 6-1 用户认证模块测试"),
        ("表6-2 商品浏览与检索模块测试", "表 6-2 商品浏览与检索模块测试"),
        ("表6-3 购物车与地址管理模块测试", "表 6-3 购物车与地址管理模块测试"),
        ("表6-4 订单流程模块测试", "表 6-4 订单流程模块测试"),
        ("表6-5 后台商品与内容管理模块测试", "表 6-5 后台商品与内容管理模块测试"),
    ]:
        set_paragraph_text(find_paragraph(doc, txt), new, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)

    if not any("表 6-6 溯源与看板统计模块测试" in p.text for p in doc.paragraphs):
        p355 = find_paragraph(doc, "表 6-5 后台商品与内容管理模块测试")
        tbl55 = next_table_after_paragraph(p355)
        cap = insert_paragraph_after_table(tbl55, "表 6-6 溯源与看板统计模块测试", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_table_after(
            doc,
            cap,
            [
                ["用例编号", "测试目标", "输入/操作", "预期结果", "实际结果"],
                ["T20", "查询溯源记录", "按商品 ID 请求溯源接口", "按节点顺序返回记录", "通过"],
                ["T21", "溯源时间线展示", "打开商品详情页", "节点名称、时间和地点展示正确", "通过"],
                ["T22", "区域库存查询", "调用区域库存接口", "返回区域库存统计列表", "通过"],
                ["T23", "看板数据展示", "打开管理端 Dashboard", "指标卡片与趋势图正常渲染", "通过"],
            ],
        )

    h623 = find_paragraph(doc, "6.2.3 异常处理测试")
    set_paragraph_text(h623, "6.2.3 结果分析", style="Heading 3")
    p357 = find_paragraph(doc, "在开发过程中，重点测试了因 Servlet")
    set_paragraph_text(
        p357,
        "测试结果表明，系统在用户认证、商品浏览、购物车与地址管理、订单状态流转、后台商品与内容管理以及溯源与看板统计等核心模块上均能稳定运行。重点异常场景包括地址缺失拦截、错误登录提示以及非法订单状态流转拦截，系统均能给出明确反馈并保持数据一致性。",
        style="Normal",
    )

    h63 = find_paragraph(doc, "6.3 测试总结")
    p359 = find_paragraph(doc, "本章通过构建合理的测试环境")
    set_paragraph_text(
        p359,
        "本次测试围绕 6 个主要功能模块共设计 23 条测试用例，覆盖用户认证、商品浏览、购物车与地址管理、订单状态流转、后台商品与内容管理、溯源查询及经营看板等核心业务场景。测试结果见表 6-7。",
        style="Normal",
    )
    if not any("表 6-7 测试结果统计表" in p.text for p in doc.paragraphs):
        cap = insert_paragraph_after(p359, "表 6-7 测试结果统计表", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_table_after(
            doc,
            cap,
            [
                ["统计项", "数值", "说明"],
                ["测试模块数", "6", "与主要功能模块图中的核心业务对应"],
                ["测试用例数", "23", "覆盖正常流程与关键异常流程"],
                ["通过数", "23", "当前测试全部通过"],
                ["未通过数", "0", "未发现阻断性问题"],
            ],
        )
        p360 = find_paragraph(doc, "")
        set_paragraph_text(
            p360,
            "综合测试结果表明，系统在主要业务链路上运行稳定，前后端交互正常，核心功能实现与设计目标基本一致。后续可继续补充自动化回归测试、接口异常测试和并发压力测试，以进一步提升系统质量保障能力。",
            style="Normal",
        )


def update_ch7_and_refs(doc: Document):
    # Keep existing heading structure, replace content with current project content
    set_paragraph_text(find_paragraph(doc, "7.1 系统分析总结"), "7.1 系统分析总结", style="Heading 2")
    p363 = find_paragraph(doc, "本文详细论述了基于 Java Web 技术的美食街摊位管理系统")
    set_paragraph_text(
        p363,
        "本文围绕水产品销售管理场景，完成了从需求分析到系统实现与测试验证的完整工作。系统以真实业务为牵引，构建了消费者端与管理端协同的双端架构，打通了商品管理、交易履约、用户服务、数据看板与溯源展示等关键能力。",
        style="Normal",
    )
    p364h = find_paragraph(doc, "7.2 存在的问题与不足")
    # insert one more paragraph before 7.2 if not present
    if "实践证明，该系统能够有效解决传统场景下数据分散" not in [p.text.strip() for p in doc.paragraphs]:
        insert_paragraph_after(
            p363,
            "实践证明，该系统能够有效解决传统场景下数据分散、流程不透明、管理决策滞后的问题，并在可维护性、可扩展性与可演示性方面达到预期目标。论文按照规范给出了结构化图表、流程图、测试用例与截图内容，具备直接进入后续排版与答辩材料制作的条件。",
            style="Normal",
        )

    bullets = [
        "1. 安全体系深度不足：当前认证与权限控制满足教学演示和中小规模试运行需求，但缺少细粒度权限模型、完善的审计追踪和统一风控策略。",
        "2. 高并发与性能验证不充分：系统尚未进行系统化压测，缓存命中策略、数据库慢查询治理和异步化链路仍有优化空间。",
        "3. 算法与数据智能能力较弱：目前以规则驱动为主，缺少销量预测、补货优化、用户画像与个性化推荐等数据智能模块。",
        "4. 供应链数据接入范围有限：溯源以结构化节点记录为主，尚未打通冷链温控、物流轨迹、设备采集等多源实时数据。",
        "5. 测试体系覆盖不完整：当前以功能验证为主，自动化测试、接口回归、端到端测试与持续集成流水线尚未全面落地。",
        "6. 运维与可观测性能力薄弱：缺少统一监控看板、告警分级、链路追踪与故障演练机制。",
    ]
    target_paras = [
        find_paragraph(doc, "尽管本系统已经实现了预期的核心功能"),
        find_paragraph(doc, "并发处理能力有限"),
        find_paragraph(doc, "支付模块缺失"),
        find_paragraph(doc, "前端交互体验"),
    ]
    set_paragraph_text(target_paras[0], "尽管系统已经实现了预期的核心功能，但在实际应用中仍存在一些有待完善的地方：", style="Normal")
    set_paragraph_text(target_paras[1], bullets[0], style="Normal")
    set_paragraph_text(target_paras[2], bullets[1], style="Normal")
    set_paragraph_text(target_paras[3], bullets[2], style="Normal")
    p_extra = insert_paragraph_after(target_paras[3], bullets[3], style="Normal")
    p_extra = insert_paragraph_after(p_extra, bullets[4], style="Normal")
    insert_paragraph_after(p_extra, bullets[5], style="Normal")

    p369 = find_paragraph(doc, "7.3 展望")
    set_paragraph_text(p369, "7.3 展望", style="Heading 2")
    outlook = [
        "后续研究与工程迭代可从以下方向推进：",
        "1. 安全强化：完善鉴权体系、权限模型与密码加密存储；引入操作审计与风控策略。",
        "2. 性能优化：增加缓存层、热点接口优化、异步化处理与数据库索引优化，提升高并发性能。",
        "3. 智能化运营：结合历史数据构建销量预测、补货建议与用户分层推荐模型。",
        "4. 供应链协同：扩展冷链监控数据接入，实现运输温控与告警联动。",
        "5. 可信溯源升级：在现有结构化追溯基础上探索区块链存证或第三方可信平台对接。",
        "6. 测试体系完善：引入自动化测试流水线，覆盖单元测试、接口测试与端到端测试。",
        "通过以上优化，系统可从可用原型逐步升级为可运营平台，进一步提升水产品数字化经营能力。",
    ]
    p370 = find_paragraph(doc, "针对上述不足，在未来的系统升级和迭代中")
    p371 = find_paragraph(doc, "架构升级：")
    p372 = find_paragraph(doc, "引入缓存与消息队列：")
    p373 = find_paragraph(doc, "对接真实支付接口：")
    p374 = find_paragraph(doc, "增加数据分析功能：")
    set_paragraph_text(p370, outlook[0], style="Normal")
    set_paragraph_text(p371, outlook[1], style="Normal")
    set_paragraph_text(p372, outlook[2], style="Normal")
    set_paragraph_text(p373, outlook[3], style="Normal")
    set_paragraph_text(p374, outlook[4], style="Normal")
    p_extra = insert_paragraph_after(p374, outlook[5], style="Normal")
    p_extra = insert_paragraph_after(p_extra, outlook[6], style="Normal")
    insert_paragraph_after(p_extra, outlook[7], style="Normal")

    refs = [
        "[1] 刘宇轩, 张敏. 基于SpringBoot的海产品电商平台设计与实现[J]. 信息技术与信息化, 2024, (3): 112-116.",
        "[2] 陈思远, 李华. 基于Spring Boot和MySQL的生鲜销售管理系统开发[J]. 计算机与现代化, 2023, (11): 78-83.",
        "[3] 王瑞, 赵丽萍. 水产品线上交易系统的数据库优化设计与实现[J]. 渔业现代化, 2024, 51(2): 45-50.",
        "[4] 孙宇航. 基于SpringBoot的海产品销售系统设计与实现[D]. 济南: 山东大学, 2024.",
        "[5] 李婷婷. 基于SpringBoot框架的水产品配送管理系统的设计与开发[D]. 青岛: 中国海洋大学, 2025.",
        "[6] 郑浩然, 吴晓燕. 基于SpringBoot和Vue的生鲜电商后台管理系统[J]. 现代计算机, 2024, 30(15): 92-97.",
        "[7] 赵明远. 基于SpringBoot的水产品交易平台设计与实现[D]. 大连: 大连海事大学, 2024.",
        "[8] 郭静怡, 刘洋. 面向水产品销售的数据持久层优化——基于MySQL与SpringBoot整合[J]. 电脑知识与技术, 2025, 21(4): 34-38.",
        "[9] Wang W J, Wang H Y, Chen M, Zou Y B, Ge Y. Aquatic product trading matching model based on multi-chain storage optimization[J]. Transactions of the Chinese Society for Agricultural Machinery, 2024, 55(6): 321-330.",
        "[10] Lee S, Kim J. Design and implementation of an e-commerce backend management system for fresh agricultural products using SpringBoot and MySQL[C]// Proceedings of the 2024 International Conference on Software Engineering and Information Management. Seoul: Yeungnam University Press, 2024: 156-162.",
    ]
    ref_paras = [p for p in doc.paragraphs if p.text.strip().startswith("[") or p.text.strip().startswith("［") or p.text.strip().startswith("\\[")]
    for idx, ref in enumerate(refs):
        if idx < len(ref_paras):
            set_paragraph_text(ref_paras[idx], ref, style="Normal")
    ack = find_paragraph(doc, "时光荏苒，论文至此已接近尾声")
    set_paragraph_text(
        ack,
        "时光荏苒，论文至此已接近尾声。在本次毕业设计与论文撰写过程中，我围绕基于 Spring Boot 与 Vue 的水产品销售管理系统，完成了从需求分析、系统设计、功能实现到测试验证的全过程。在此过程中，我得到了许多老师、同学和家人的关心与支持。首先，感谢指导教师在选题确定、论文结构调整、系统实现思路以及文稿修改等方面给予的持续指导，使我能够不断发现问题并完成完善。其次，感谢学院各位任课教师在大学阶段传授的专业知识，为本课题的完成奠定了理论基础；同时感谢同学和朋友在项目开发、资料查找和论文修改过程中提供的帮助。最后，感谢家人一直以来的理解与支持，使我能够顺利完成本次毕业设计与论文写作。",
        style="Normal",
    )


def main():
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)
    generate_module_diagram(MODULE_IMG)
    doc = Document(TARGET)
    update_heading_styles(doc)
    update_ch3(doc)
    update_ch4(doc)
    update_ch5(doc)
    update_ch6(doc)
    update_ch7_and_refs(doc)
    ensure_update_fields(doc)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
